# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
FIG = OUT / "figures"
REPORT = OUT / "20255227021_朱旭东_机器学习课程考核报告.docx"
COVER = OUT / "20255227021_朱旭东_机器学习课程考核封面.docx"
ZIP_PATH = OUT / "20255227021_朱旭东_机器学习课程考核提交包.zip"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.space_after = Pt(4)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 3"].font.size = Pt(10.5)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("基于 LSTM、Transformer 与 CNN-Transformer 的家庭电力消耗预测研究")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("2026 年专硕机器学习课程考核报告")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.first_line_indent = Pt(0)
    run = meta.add_run("姓名：朱旭东    学号：20255227021    完成方式：单人完成")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def add_para(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)
        rest = text[len(bold_prefix) :]
        if rest:
            p.add_run(rest)
    else:
        p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.first_line_indent = None
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.first_line_indent = None
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=8.5)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.5)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
                row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def add_picture(doc: Document, path: Path, caption: str, width: float = 5.9) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)


def fmt(x: str | float, digits: int = 4) -> str:
    return f"{float(x):.{digits}f}"


def fmt_k(x: str | float, digits: int = 2) -> str:
    return f"{float(x) / 1000.0:.{digits}f}"


def build_report() -> None:
    summary_path = OUT / "summary.csv"
    results_path = OUT / "results.csv"
    meta_path = OUT / "metadata.json"
    best_path = OUT / "best_plots.csv"
    if not summary_path.exists() or not results_path.exists() or not meta_path.exists():
        raise FileNotFoundError("Please run run_experiment.py before make_report.py.")

    summary = read_csv(summary_path)
    results = read_csv(results_path)
    best = read_csv(best_path)
    metadata = json.loads(meta_path.read_text(encoding="utf-8"))

    doc = Document()
    set_doc_styles(doc)
    add_page_number(doc.sections[0])
    add_title(doc)

    doc.add_heading("摘要", level=1)
    add_para(
        doc,
        "本文围绕家庭电力消耗预测问题，构建多变量时间序列建模流程，使用过去 90 天的用电与日历特征预测未来 90 天和 365 天的 global_active_power。实验比较了 LSTM、Transformer 以及本文提出的 CNN-Transformer 改进模型。每种模型在两种预测长度上分别进行 5 轮随机种子实验，并以 MSE 和 MAE 的均值、标准差评价预测精度与稳定性。结果显示，在当前数据规模和轻量训练设置下，LSTM 取得了最好的平均误差；CNN-Transformer 的主要价值在于为开放题引入局部卷积归纳偏置，并在长期预测中表现出比纯 Transformer 更低的误差。"
    )

    doc.add_heading("1. 问题介绍", level=1)
    add_para(
        doc,
        "随着智能家居和智能电网的发展，家庭电力消耗预测不仅能帮助居民识别异常用电、安排峰谷用电，还能为电网调度和负荷管理提供参考。家庭用电受季节、日期、家庭行为模式和外部气象等因素影响，呈现明显的非线性和多尺度时间依赖，因此适合作为机器学习课程中时间序列建模的综合实验。"
    )
    add_para(
        doc,
        f"本实验使用课程要求指定的 Individual household electric power consumption 数据。当前复现实验的数据源为：{metadata.get('data_source')}。数据时间范围为 {metadata['start_date']} 至 {metadata['end_date']}，共 {metadata['n_days']} 个日尺度样本。由于本地未提供单独的 train.csv 与 test/tes.csv，代码默认从 UCI 数据集构建日尺度序列；若后续获得课程发布的数据文件，只需放入 data 目录即可复跑同一流程。"
    )
    add_para(
        doc,
        "原始分钟级数据按课程说明进行日汇总：global_active_power、global_reactive_power 与各 sub_metering 字段取日总和，voltage 和 global_intensity 取日均值，并计算 sub_metering_remainder。为了增强模型对周期变化的感知，本文额外加入月份、星期、年内日序以及周末标记等日历变量。"
    )
    add_picture(doc, FIG / "daily_power_overview.png", "图1 日尺度 global_active_power 序列概览")

    doc.add_heading("2. 模型", level=1)
    add_para(
        doc,
        "设第 t 天的多变量特征为 x_t，模型输入为最近 90 天序列 X=[x_{t-89},...,x_t]，输出为未来 H 天的总有功功率 y=[p_{t+1},...,p_{t+H}]，其中 H 分别取 90 和 365。两个预测长度分别训练模型，避免用长期预测模型参数替代短期预测模型。"
    )
    doc.add_heading("2.1 LSTM 基线模型", level=2)
    add_para(
        doc,
        "LSTM 通过输入门、遗忘门和输出门缓解普通循环网络的梯度衰减问题，适合建模家庭用电中的惯性和短中期依赖。本文使用单层 LSTM 编码 90 天历史序列，并将最后一个隐状态输入全连接层，一次性输出未来 H 天预测值。"
    )
    doc.add_heading("2.2 Transformer 模型", level=2)
    add_para(
        doc,
        "Transformer 使用自注意力机制直接建立任意时间位置之间的依赖关系。本文先将每日多变量特征线性映射到 d_model 维空间，加入正弦位置编码后输入 Transformer Encoder，再对时间维进行平均池化，最后由线性层输出 H 天预测结果。该模型对长期依赖更友好，但在样本量较小时也更容易受到噪声和过拟合影响。"
    )
    doc.add_heading("2.3 改进模型 CNN-Transformer", level=2)
    add_para(
        doc,
        "本文提出的改进模型在 Transformer 前加入一维卷积模块。卷积层在时间轴上提取局部趋势、周内波动和短期突变模式，再由 Transformer 建模较长范围的依赖关系。其动机是家庭用电同时包含局部周期性与长期季节性，单纯自注意力对小数据集可能缺少局部归纳偏置，而卷积可以为 Transformer 提供更平滑、更稳定的局部表征。"
    )
    add_para(
        doc,
        "简化伪代码如下：1）读取并按天汇总数据；2）用过去 90 天窗口构造样本；3）对每个 horizon 和随机种子分别训练 LSTM、Transformer、CNN-Transformer；4）在测试集上计算 MSE、MAE；5）汇总 5 轮均值和标准差并绘制预测曲线。"
    )

    doc.add_heading("3. 结果与分析", level=1)
    add_para(
        doc,
        "实验采用按时间顺序切分的训练集、验证集和测试集，避免未来信息泄漏。评价指标为 MSE 与 MAE。MSE 对大误差更敏感，适合衡量预测曲线在峰值处的偏离；MAE 具有更直观的平均绝对偏差含义。"
    )
    doc.add_page_break()
    rows = []
    model_label = {"ConvTransformer": "CNN-Trans", "Transformer": "Trans", "LSTM": "LSTM"}
    for r in summary:
        rows.append(
            [
                model_label.get(r["model"], r["model"]),
                f"{r['horizon']} 天",
                fmt_k(r["mse_mean"]),
                fmt_k(r["mse_std"]),
                fmt(r["mae_mean"], 2),
                fmt(r["mae_std"], 2),
                fmt(r["epochs_mean"], 2),
            ]
        )
    add_para(doc, "表中 MSE 均值和标准差按 10^3 缩放显示，以避免数值过长影响版式。")
    add_table(
        doc,
        ["模型", "长度", "MSE均值\n(×10^3)", "MSE std\n(×10^3)", "MAE均值", "MAE std", "轮数"],
        rows,
        widths=[0.90, 0.65, 1.05, 1.05, 0.85, 0.80, 0.80],
    )

    baselines = metadata.get("horizon_90_baseline", {}), metadata.get("horizon_365_baseline", {})
    add_para(
        doc,
        f"作为参照，持久性基线在 90 天预测上的 MSE/MAE 为 {baselines[0].get('mse', 0):.4f}/{baselines[0].get('mae', 0):.4f}，在 365 天预测上的 MSE/MAE 为 {baselines[1].get('mse', 0):.4f}/{baselines[1].get('mae', 0):.4f}。与该简单基线相比，神经网络模型能够利用多变量历史窗口拟合更复杂的非线性关系。"
    )

    for horizon in [90, 365]:
        doc.add_heading(f"3.{1 if horizon == 90 else 2} {horizon} 天预测曲线", level=2)
        subset = [b for b in best if int(b["horizon"]) == horizon]
        for b in subset:
            path = Path(b["plot_path"])
            add_picture(
                doc,
                path,
                f"图 {horizon}-{b['model']} 预测曲线与 Ground Truth 对比（seed={b['seed']}）",
                width=5.85,
            )

    add_para(
        doc,
        "从整体趋势看，90 天预测更容易保持相位和幅值一致，365 天预测则会出现更明显的均值回归和峰值平滑现象。这说明模型虽然能够学习长期季节趋势，但对一年尺度内的局部异常、节假日行为变化和未观测外部因素仍较敏感。"
    )
    add_para(
        doc,
        "三类模型的差异主要体现在归纳偏置上。LSTM 按时间递推，结构简单且参数量较小，在本实验中反而更适合 1400 多天规模的日尺度数据；Transformer 能直接比较所有历史位置，对长期依赖更自然，但样本量有限时容易过度依赖噪声；CNN-Transformer 在输入端加入局部卷积，使模型先捕获邻近天数的平滑模式，再进行全局注意力建模。该改进模型没有在短期预测上超过 LSTM，但在 365 天预测中优于纯 Transformer，说明局部卷积对长期预测的稳定性有一定帮助。"
    )

    doc.add_heading("4. 讨论", level=1)
    add_para(
        doc,
        "本实验的主要限制包括三点。第一，本地未获得课程单独发布的 train.csv 与 tes.csv，因此实验默认使用课程说明中的 UCI 原始数据自行构建日尺度样本；代码已经保留对课程 CSV 的自动读取接口。第二，天气数据在 PDF 中作为可融合变量给出，但开放站点的数据需要进一步确定测站和月份映射，本次报告以电力变量和日历变量为主，后续可将 RR、NBJRR1、NBJRR5、NBJRR10 与 NBJBROU 等字段并入同一特征表。第三，为保证 CPU 环境可在合理时间内完成 30 组实验，模型规模和训练轮数较为克制，绝对性能仍有提升空间。"
    )
    add_para(
        doc,
        "后续可从三个方向改进：一是加入更完整的天气、节假日和温度变量，增强模型对外部驱动因素的解释能力；二是采用直接多步预测与递归预测相结合的策略，缓解长预测 horizon 下的均值化问题；三是尝试 PatchTST、Informer 或多尺度分解模型，将年周期、周周期和短期扰动分开建模。"
    )
    add_para(
        doc,
        "工具使用说明：报告文字整理过程中使用了 ChatGPT/Codex 辅助完成表述润色、代码组织与文档生成；实验代码、指标、图表均由本项目脚本在本地生成。"
    )

    doc.add_heading("参考文献", level=1)
    refs = [
        "Dua, D. and Graff, C. UCI Machine Learning Repository: Individual household electric power consumption Data Set.",
        "Hochreiter, S. and Schmidhuber, J. Long Short-Term Memory. Neural Computation, 1997.",
        "Vaswani, A. et al. Attention Is All You Need. NeurIPS, 2017.",
        "Lim, B. and Zohren, S. Time-series forecasting with deep learning: a survey. Philosophical Transactions of the Royal Society A, 2021.",
        "课程考核 PDF：2026 年专硕机器学习课程项目。",
    ]
    for ref in refs:
        add_number(doc, ref)

    doc.add_heading("附录：代码与贡献说明", level=1)
    add_para(doc, "代码链接：本地 Git 仓库已在 ml_power_forecast_final 目录初始化并提交。由于当前环境没有 GitHub CLI 和登录凭据，提交前需将该目录推送到个人 GitHub 仓库，并把仓库 URL 填写到此处。")
    add_para(doc, "单人完成，朱旭东负责数据处理、模型实现、实验运行、结果分析和报告撰写。")
    add_para(doc, f"实验脚本输出文件包括 results.csv、summary.csv、metadata.json 和 figures 目录。单次结果共 {len(results)} 条，覆盖 3 个模型 × 2 个预测长度 × 5 个随机种子。")

    doc.save(REPORT)


def build_cover() -> None:
    doc = Document()
    set_doc_styles(doc)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("机器学习课程考核报告")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(24)
    for _ in range(3):
        doc.add_paragraph()
    fields = [
        ("题目", "基于 LSTM、Transformer 与 CNN-Transformer 的家庭电力消耗预测研究"),
        ("姓名", "朱旭东"),
        ("学号", "20255227021"),
        ("完成方式", "单人完成"),
        ("课程", "2026 年专硕机器学习"),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (k, v) in enumerate(fields):
        set_cell_text(table.rows[i].cells[0], k, bold=True, size=11)
        set_cell_text(table.rows[i].cells[1], v, size=11)
        table.rows[i].cells[0].width = Inches(1.6)
        table.rows[i].cells[1].width = Inches(4.8)
    doc.save(COVER)


def build_zip() -> None:
    include_files = [
        REPORT,
        COVER,
        OUT / "results.csv",
        OUT / "summary.csv",
        OUT / "metadata.json",
        ROOT / "run_experiment.py",
        ROOT / "make_report.py",
        ROOT / "set_github_url.py",
        ROOT / "README.md",
        ROOT / "requirements.txt",
        ROOT / "exam_requirements_extracted.txt",
    ]
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in include_files:
            if path.exists():
                zf.write(path, arcname=path.relative_to(ROOT))
        for fig in sorted(FIG.glob("*.png")):
            zf.write(fig, arcname=fig.relative_to(ROOT))


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_report()
    build_cover()
    build_zip()
    print(REPORT)
    print(COVER)
    print(ZIP_PATH)


if __name__ == "__main__":
    main()
