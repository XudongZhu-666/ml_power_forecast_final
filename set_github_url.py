# -*- coding: utf-8 -*-
from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
REPORT = OUT / "20255227021_朱旭东_机器学习课程考核报告.docx"
ZIP_PATH = OUT / "20255227021_朱旭东_机器学习课程考核提交包.zip"


def update_report(url: str) -> None:
    doc = Document(str(REPORT))
    replaced = False
    for para in doc.paragraphs:
        if para.text.startswith("代码链接："):
            para.text = f"代码链接：{url}"
            replaced = True
            break
    if not replaced:
        doc.add_paragraph(f"代码链接：{url}")
    doc.save(REPORT)


def update_readme(url: str) -> None:
    readme = ROOT / "README.md"
    text = readme.read_text(encoding="utf-8")
    marker = "## 代码链接"
    replacement = f"## 代码链接\n\n{url}\n"
    if marker in text:
        head = text.split(marker)[0].rstrip()
        text = head + "\n\n" + replacement
    else:
        text = text.rstrip() + "\n\n" + replacement
    readme.write_text(text, encoding="utf-8")


def rebuild_zip() -> None:
    include_files = [
        REPORT,
        OUT / "20255227021_朱旭东_机器学习课程考核封面.docx",
        OUT / "results.csv",
        OUT / "summary.csv",
        OUT / "metadata.json",
        ROOT / "run_experiment.py",
        ROOT / "make_report.py",
        ROOT / "set_github_url.py",
        ROOT / "README.md",
        ROOT / "requirements.txt",
        ROOT / "exam_requirements_extracted.txt",
    ]
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in include_files:
            if path.exists():
                zf.write(path, arcname=path.relative_to(ROOT))
        for fig in sorted((OUT / "figures").glob("*.png")):
            zf.write(fig, arcname=fig.relative_to(ROOT))


def main() -> None:
    if len(sys.argv) != 2 or not sys.argv[1].startswith(("https://github.com/", "http://github.com/")):
        raise SystemExit("Usage: python set_github_url.py https://github.com/<user>/<repo>")
    url = sys.argv[1].strip()
    update_report(url)
    update_readme(url)
    rebuild_zip()
    print(f"Updated GitHub URL: {url}")
    print(REPORT)
    print(ZIP_PATH)


if __name__ == "__main__":
    main()
